"""
Career Service — phân tích CV và gợi ý nghề nghiệp.
Rule-based mặc định; nâng AI khi có api_key.
"""

from __future__ import annotations

import base64
import io
import logging
import re
import tempfile
from pathlib import Path

from models.schemas import (
    EvaluateCVRequest,
    EvaluateCVResponse,
    ParseCVRequest,
    ParseCVResponse,
    RecommendRequest,
    RecommendResponse,
)
from services.provider import call_provider
from utils.json_call import complete_json

logger = logging.getLogger(__name__)


# =============================================================================
# Bản đồ kỹ năng theo vai trò
# =============================================================================

ROLE_SKILL_MAP: dict[str, list[str]] = {
    # --- Công nghệ thông tin ---
    "laravel": ["PHP", "Laravel", "MySQL", "REST API", "Git", "Docker"],
    "php": ["PHP", "Laravel", "MySQL", "REST API", "Git", "Docker"],
    "backend": ["PHP", "Laravel", "MySQL", "REST API", "Testing", "Docker"],
    "frontend": ["JavaScript", "TypeScript", "Vue.js", "Nuxt", "HTML/CSS", "Git"],
    "full stack": ["PHP", "Laravel", "JavaScript", "Vue.js", "MySQL", "Docker"],
    "fullstack": ["PHP", "Laravel", "JavaScript", "Vue.js", "MySQL", "Docker"],
    "devops": ["Docker", "Linux", "AWS", "Git", "CI/CD", "Testing"],
    "mobile": ["Flutter", "Dart", "React Native", "JavaScript", "Git"],
    "python": ["Python", "Django", "FastAPI", "SQL", "Docker", "Git"],
    "java": ["Java", "Spring Boot", "MySQL", "REST API", "Git", "Docker"],
    "react": ["React", "JavaScript", "TypeScript", "HTML/CSS", "Git", "Node.js"],
    "vue": ["Vue.js", "Nuxt", "JavaScript", "TypeScript", "HTML/CSS", "Git"],
    "ai": ["Python", "Machine Learning", "Deep Learning", "TensorFlow", "SQL"],
    "cloud": ["AWS", "Docker", "Kubernetes", "Linux", "CI/CD", "Terraform"],
    "qa": ["Testing", "Selenium", "Postman", "SQL", "Git", "CI/CD"],
    "tester": ["Testing", "Selenium", "Postman", "SQL", "Git", "CI/CD"],
    "kiểm thử": ["Testing", "Selenium", "Postman", "SQL", "Git", "CI/CD"],
    "data scientist": ["Python", "Machine Learning", "Pandas", "SQL", "Statistics", "Data Visualization"],
    "data science": ["Python", "Machine Learning", "Pandas", "SQL", "Statistics", "Data Visualization"],
    "data engineer": ["Python", "SQL", "ETL", "Spark", "Airflow", "Data Warehouse"],
    "data": ["Python", "SQL", "Machine Learning", "Testing", "Statistics"],
    "security": ["Network Security", "Penetration Testing", "SIEM", "Linux", "Cryptography", "OWASP"],
    "cyber": ["Network Security", "Penetration Testing", "SIEM", "Linux", "Cryptography", "OWASP"],
    "an ninh mạng": ["Network Security", "Penetration Testing", "SIEM", "Linux", "Cryptography", "OWASP"],
    ".net": [".NET", "C#", "SQL Server", "ASP.NET", "Git", "REST API"],
    "c#": [".NET", "C#", "SQL Server", "ASP.NET", "Git", "REST API"],
    "golang": ["Go", "REST API", "Docker", "Git", "Microservices", "SQL"],
    "go developer": ["Go", "REST API", "Docker", "Git", "Microservices", "SQL"],
    "blockchain": ["Solidity", "Blockchain", "Smart Contract", "Web3.js", "Ethereum", "Git"],
    "game": ["Unity", "C#", "Game Design", "3D Modeling", "Git", "C++"],
    # --- Quản trị kinh doanh ---
    "product manager": ["Agile", "Scrum", "Product Roadmap", "User Story", "Analytics", "Communication"],
    "product owner": ["Agile", "Scrum", "Product Roadmap", "User Story", "Analytics", "Communication"],
    "project manager": ["Agile", "Scrum", "Project Planning", "Risk Management", "Communication", "MS Project"],
    "business analyst": ["Requirement Analysis", "SQL", "BPMN", "Documentation", "Communication", "Excel"],
    "digital marketing": ["SEO", "Google Ads", "Content Marketing", "Analytics", "Social Media", "Email Marketing"],
    "marketing": ["SEO", "Google Ads", "Content Marketing", "Analytics", "Social Media", "Email Marketing"],
    # --- Điện tử viễn thông ---
    "network": ["Networking", "Cisco", "TCP/IP", "Linux", "Security", "Firewall"],
    "mạng": ["Networking", "Cisco", "TCP/IP", "Linux", "Security", "Firewall"],
    "embedded": ["C/C++", "Embedded C", "IoT Protocols", "Microcontroller", "RTOS", "Circuit Design"],
    "nhúng": ["C/C++", "Embedded C", "IoT Protocols", "Microcontroller", "RTOS", "Circuit Design"],
    "iot": ["C/C++", "Embedded C", "IoT Protocols", "Microcontroller", "RTOS", "Circuit Design"],
}

SKILL_DICTIONARY: dict[str, list[str]] = {
    "PHP": ["php"],
    "Laravel": ["laravel"],
    "MySQL": ["mysql"],
    "PostgreSQL": ["postgresql", "postgres"],
    "MongoDB": ["mongodb", "mongo"],
    "Redis": ["redis"],
    "REST API": ["rest api", "restful"],
    "JavaScript": ["javascript"],
    "TypeScript": ["typescript"],
    "Vue.js": ["vue.js", "vuejs", "vue"],
    "Nuxt": ["nuxt"],
    "React": ["react", "reactjs"],
    "Node.js": ["node.js", "nodejs"],
    "HTML/CSS": ["html", "css", "tailwind", "bootstrap"],
    "Docker": ["docker"],
    "Git": ["git", "github", "gitlab"],
    "Linux": ["linux", "ubuntu"],
    "Java": ["java", "spring boot", "spring"],
    "Python": ["python", "django", "flask", "fastapi"],
    "C/C++": ["c++", "c/c++"],
    "Flutter": ["flutter", "dart"],
    "AWS": ["aws"],
    "Testing": ["phpunit", "unit test", "testing", "kiểm thử"],
    "SQL": ["sql"],
    "Machine Learning": ["machine learning", "deep learning"],
    "Communication": ["giao tiếp", "communication", "teamwork", "làm việc nhóm"],
}


async def parse_cv(payload: ParseCVRequest) -> ParseCVResponse:
    """Trích text từ PDF/DOCX (pdfminer) + OCR AI nếu PDF scan."""
    raw = _load_bytes(payload)
    if not raw:
        return ParseCVResponse(text="", skills=[], method="none")

    filename = (payload.file_name or payload.file_path or "cv.pdf").lower()
    ext = Path(filename).suffix.lstrip(".").lower()

    text = ""
    method = "none"
    ocr_used = False

    if ext in {"pdf", ""}:
        text, method = _extract_pdf_text(raw)
        if _text_quality(text) < 40 and payload.api_key:
            ocr_text = await _ocr_pdf_with_ai(raw, payload)
            if _text_quality(ocr_text) > _text_quality(text):
                text = ocr_text
                method = "ai_ocr"
                ocr_used = True
    elif ext == "docx":
        text, method = _extract_docx_text(raw)
    elif ext == "doc":
        text, method = "", "unsupported_doc"

    text = _normalize(text)
    skills = _extract_skills(text)
    return ParseCVResponse(
        text=text,
        skills=skills,
        method=method,
        ocr_used=ocr_used,
    )


def _load_bytes(payload: ParseCVRequest) -> bytes:
    if payload.file_base64:
        data = payload.file_base64
        if "," in data and data.strip().startswith("data:"):
            data = data.split(",", 1)[1]
        try:
            return base64.b64decode(data)
        except Exception as exc:
            logger.warning("Invalid CV base64: %s", exc)
            return b""

    if payload.file_path:
        path = Path(payload.file_path)
        if path.is_file():
            return path.read_bytes()

    return b""


def _extract_pdf_text(raw: bytes) -> tuple[str, str]:
    # pdfminer thường tốt hơn với PDF TopCV/Unicode
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp:
            tmp.write(raw)
            tmp.flush()
            text = pdfminer_extract(tmp.name) or ""
        text = _normalize(text)
        if _text_quality(text) >= 40:
            return text, "pdfminer"
    except Exception as exc:
        logger.warning("pdfminer failed: %s", exc)

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        chunks = [(page.extract_text() or "") for page in reader.pages]
        text = _normalize("\n".join(chunks))
        if text:
            return text, "pypdf"
    except Exception as exc:
        logger.warning("pypdf failed: %s", exc)

    return "", "none"


def _extract_docx_text(raw: bytes) -> tuple[str, str]:
    try:
        from docx import Document

        doc = Document(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return _normalize("\n".join(parts)), "python_docx"
    except Exception as exc:
        logger.warning("docx extract failed: %s", exc)
        return "", "none"


async def _ocr_pdf_with_ai(raw: bytes, payload: ParseCVRequest) -> str:
    """PDF ảnh/scan → lấy ảnh tối đa 5 trang đầu → Gemini/Claude/OpenAI-compat OCR trong 1 lần gọi."""
    images = _pdf_images(raw, max_pages=5)
    if not images or not payload.api_key:
        return ""

    page_note = "Đây là ảnh CV" + (f" gồm {len(images)} trang liên tiếp" if len(images) > 1 else "") + "."
    prompt = (
        f"{page_note} Hãy trích TOÀN BỘ chữ đọc được (tiếng Việt/Anh) theo đúng thứ tự trang, "
        "giữ cấu trúc theo dòng. Chỉ trả về nội dung CV, không giải thích."
    )
    provider = (payload.provider or "gemini").strip().lower()

    try:
        if provider == "gemini":
            reply, _ = await _gemini_vision(payload.api_key or "", images, prompt, payload.model)
        elif provider == "claude":
            reply, _ = await _claude_vision(payload.api_key or "", images, prompt, payload.model)
        else:
            # OpenAI-compatible multimodal — nhiều ảnh trong cùng 1 message
            content: list[dict] = [{"type": "text", "text": prompt}]
            for image_b64, mime in images:
                content.append({
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime};base64,{image_b64}"},
                })
            messages = [{"role": "user", "content": content}]
            reply, _ = await call_provider(
                provider=provider,
                api_key=payload.api_key or "",
                messages=messages,  # type: ignore[arg-type]
                model=payload.model,
                temperature=0.1,
            )
        return _normalize(reply or "")
    except Exception as exc:
        logger.warning("AI OCR failed: %s", exc)
        return ""


def _pdf_images(raw: bytes, max_pages: int = 5) -> list[tuple[str, str]]:
    """Lấy ảnh của tối đa max_pages trang đầu (mỗi trang lấy ảnh lớn nhất nếu có)."""
    images: list[tuple[str, str]] = []
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(raw))
        for page in reader.pages[:max_pages]:
            page_images = getattr(page, "images", None)
            if not page_images:
                continue
            img = page_images[0]
            data = img.data
            name = (img.name or "img.png").lower()
            mime = "image/jpeg" if name.endswith((".jpg", ".jpeg")) else "image/png"
            if name.endswith(".webp"):
                mime = "image/webp"
            images.append((base64.b64encode(data).decode("ascii"), mime))
    except Exception as exc:
        logger.warning("Extract PDF images failed: %s", exc)
    return images


async def _gemini_vision(
    api_key: str,
    images: list[tuple[str, str]],
    prompt: str,
    model: str | None,
) -> tuple[str | None, dict]:
    import httpx
    from config import settings

    model_name = model or settings.DEFAULT_GEMINI_MODEL
    url = f"{settings.GEMINI_API_URL}/{model_name}:generateContent?key={api_key}"
    parts: list[dict] = [{"text": prompt}]
    for image_b64, mime in images:
        parts.append({"inline_data": {"mime_type": mime, "data": image_b64}})
    body = {
        "contents": [{"role": "user", "parts": parts}],
        "generationConfig": {"temperature": 0.1},
    }
    async with httpx.AsyncClient() as client:
        response = await client.post(
            url, json=body, headers={"Content-Type": "application/json"}, timeout=90
        )
        response.raise_for_status()
        data = response.json()
    parts_out = (((data.get("candidates") or [{}])[0].get("content") or {}).get("parts")) or []
    text = "\n".join(p.get("text", "") for p in parts_out if p.get("text")).strip()
    return text or None, {}


async def _claude_vision(
    api_key: str,
    images: list[tuple[str, str]],
    prompt: str,
    model: str | None,
) -> tuple[str | None, dict]:
    import httpx
    from config import settings

    content: list[dict] = []
    for image_b64, mime in images:
        content.append({
            "type": "image",
            "source": {
                "type": "base64",
                "media_type": mime,
                "data": image_b64,
            },
        })
    content.append({"type": "text", "text": prompt})

    body = {
        "model": model or settings.DEFAULT_CLAUDE_MODEL,
        "max_tokens": 4096,
        "temperature": 0.1,
        "messages": [{"role": "user", "content": content}],
    }
    headers = {
        "Content-Type": "application/json",
        "anthropic-version": settings.CLAUDE_API_VERSION,
    }
    auth_mode = (settings.CLAUDE_AUTH_MODE or "bearer").strip().lower()
    if auth_mode == "x-api-key":
        headers["x-api-key"] = api_key
    else:
        headers["Authorization"] = f"Bearer {api_key}"

    async with httpx.AsyncClient() as client:
        response = await client.post(
            settings.CLAUDE_API_URL, json=body, headers=headers, timeout=90
        )
        response.raise_for_status()
        data = response.json()

    blocks = data.get("content") or []
    text = "\n".join(
        b.get("text", "") for b in blocks if b.get("type") == "text"
    ).strip()
    return text or None, {}


def _normalize(text: str) -> str:
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", " ", text or "")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _text_quality(text: str) -> int:
    """Đếm ký tự chữ/số có nghĩa — loại binary garbage."""
    if not text:
        return 0
    letters = re.findall(r"[A-Za-zÀ-ỹ0-9]", text)
    return len(letters)


def _extract_skills(text: str) -> list[str]:
    normalized = f" {text.lower()} "
    skills: list[str] = []
    for label, keywords in SKILL_DICTIONARY.items():
        for keyword in keywords:
            if keyword.lower() in normalized:
                skills.append(label)
                break
    return skills


async def evaluate_cv(payload: EvaluateCVRequest) -> EvaluateCVResponse:
    """Đánh giá CV theo góc nhìn nhà tuyển dụng (AI → lỗi thì 502 để Laravel fallback provider)."""
    from fastapi import HTTPException

    if payload.api_key:
        try:
            ai_result = await _evaluate_with_ai(payload)
            if ai_result is not None:
                return ai_result
            raise HTTPException(status_code=502, detail="AI evaluate returned empty")
        except HTTPException:
            raise
        except Exception as exc:
            logger.warning("Career AI evaluate failed: %s", exc)
            raise HTTPException(status_code=502, detail=f"AI evaluate failed: {exc}") from exc

    return _evaluate_rules(payload)


async def recommend(payload: RecommendRequest) -> RecommendResponse:
    """Gợi ý nghề nghiệp: ưu tiên AI khi có key, fallback rule-based."""
    if payload.api_key:
        try:
            ai_result = await _recommend_with_ai(payload)
            if ai_result is not None:
                return ai_result
        except Exception as exc:
            logger.warning("Career AI recommend failed, using rules: %s", exc)

    return _recommend_rules(payload)


async def _evaluate_with_ai(payload: EvaluateCVRequest) -> EvaluateCVResponse | None:
    skills_text = ", ".join(payload.skills) if payload.skills else "(chưa liệt kê)"
    cv_excerpt = (payload.cv_text or "")[:6000]
    salary_line = (
        f"Mức lương mong muốn: {payload.expected_salary:,} VND".replace(",", ".")
        if payload.expected_salary
        else "Chưa nêu mức lương mong muốn."
    )

    messages = [
        {
            "role": "system",
            "content": (
                "Bạn là Technical Recruiter / Hiring Manager tại công ty công nghệ Việt Nam. "
                "Đánh giá CV ứng viên một cách khách quan, thẳng thắn, cụ thể — không nịnh, không chung chung. "
                "So sánh trực tiếp với yêu cầu thực tế của vị trí mục tiêu. "
                "Trả về JSON thuần (không markdown) đúng schema:\n"
                "{\n"
                '  "match_score": 0-100,\n'
                '  "verdict": "1 câu kết luận tuyển dụng (VD: Chưa sẵn sàng / Cần chỉnh sửa / Đạt vòng CV)",\n'
                '  "overview": "đoạn nhận xét tổng quan 4-7 câu tiếng Việt, nêu cảm nhận khi đọc CV và mức phù hợp JD",\n'
                '  "strengths": ["điểm mạnh cụ thể 1", "..."] tối đa 6,\n'
                '  "weaknesses": ["điểm yếu cụ thể 1", "..."] tối đa 6,\n'
                '  "improvements": ["việc cần sửa trên CV / hồ sơ 1", "..."] tối đa 7,\n'
                '  "missing_items": ["nội dung còn thiếu so với JD 1", "..."] tối đa 6,\n'
                '  "skill_gaps": ["kỹ năng thiếu so với vị trí", "..."] tối đa 8,\n'
                '  "interview_focus": ["câu/chủ đề HR hoặc kỹ thuật sẽ hỏi nếu gọi phỏng vấn", "..."] tối đa 5,\n'
                '  "recommended_keyword_topics": ["chủ đề học để bổ sung gap", "..."] tối đa 6,\n'
                '  "salary_comment": "nhận xét ngắn về mức lương kỳ vọng so với hồ sơ, hoặc null"\n'
                "}\n"
                "match_score phải phản ánh mức khớp VỊ TRÍ MỤC TIÊU (không phải độ đầy đủ checklist). "
                "Nếu stack CV lệch JD (vd PHP/Laravel mà apply Java Fullstack) thì điểm phải thấp và nêu rõ mismatch."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Vị trí đang ứng tuyển: {payload.target_job}\n"
                f"{salary_line}\n"
                f"Kỹ năng hệ thống đã trích: {skills_text}\n\n"
                f"Nội dung CV:\n{cv_excerpt or '(không có text)'}\n\n"
                "Hãy viết bài đánh giá như bạn đang quyết định có gọi ứng viên vào vòng phỏng vấn hay không."
            ),
        },
    ]

    raw = await complete_json(
        provider=payload.provider or "gemini",
        api_key=payload.api_key or "",
        model=payload.model,
        messages=messages,
        short_messages=[
            {
                "role": "system",
                "content": (
                    'Chỉ trả JSON: {"match_score":0,"verdict":"...","overview":"...",'
                    '"strengths":[],"weaknesses":[],"improvements":[],"missing_items":[],'
                    '"skill_gaps":[],"interview_focus":[],"recommended_keyword_topics":[],'
                    '"salary_comment":null}. Không markdown.'
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Vị trí: {payload.target_job}. "
                    f"Kỹ năng: {skills_text}. "
                    f"CV (rút):\n{(payload.cv_text or '')[:1500] or '(trống)'}\n"
                    "JSON ngắn, overview 3-4 câu."
                ),
            },
        ],
        temperature=0.35,
        is_valid=lambda d: bool(str(d.get("overview") or "").strip()),
    )

    try:
        score = int(raw.get("match_score") or 0)
        score = max(0, min(score, 100))

        def _list(key: str, limit: int) -> list[str]:
            return [str(x).strip() for x in (raw.get(key) or []) if str(x).strip()][:limit]

        overview = str(raw.get("overview") or "").strip()
        verdict = str(raw.get("verdict") or "").strip()
        salary_comment = raw.get("salary_comment")
        if salary_comment is not None:
            salary_comment = str(salary_comment).strip() or None

        gaps = _list("skill_gaps", 8)
        return EvaluateCVResponse(
            match_score=score,
            verdict=verdict or "Cần chỉnh sửa hồ sơ trước khi ứng tuyển.",
            overview=overview,
            strengths=_list("strengths", 6),
            weaknesses=_list("weaknesses", 6),
            improvements=_list("improvements", 7),
            missing_items=_list("missing_items", 6),
            skill_gaps=gaps,
            interview_focus=_list("interview_focus", 5),
            recommended_keyword_topics=_list("recommended_keyword_topics", 6) or gaps[:4],
            salary_comment=salary_comment,
        )
    except (TypeError, ValueError) as exc:
        logger.warning("Failed to map evaluate-cv JSON: %s", exc)
        return None


def _evaluate_rules(payload: EvaluateCVRequest) -> EvaluateCVResponse:
    """Fallback khi không có AI: đánh giá rule dựa trên skill gap."""
    recommend = _recommend_rules(
        RecommendRequest(
            skills=payload.skills,
            target_job=payload.target_job,
            cv_text=payload.cv_text,
        )
    )
    text = (payload.cv_text or "").lower()
    skills_l = {s.lower() for s in payload.skills}
    job = (payload.target_job or "").lower()

    def _has_skill_token(*tokens: str) -> bool:
        for token in tokens:
            t = token.lower()
            if any(re.search(rf"(?<![a-z]){re.escape(t)}(?![a-z])", s) for s in skills_l):
                return True
            if re.search(rf"(?<![a-z]){re.escape(t)}(?![a-z])", text):
                return True
        return False

    mismatch_hints: list[str] = []
    if any(k in job for k in ("java", "spring")) and not _has_skill_token("java", "spring", "spring boot"):
        mismatch_hints.append(
            "CV gần như không thể hiện Java/Spring trong khi vị trí mục tiêu là Fullstack/Backend Java — "
            "đây là lệch stack nghiêm trọng với JD."
        )
    if any(k in job for k in ("php", "laravel")) and _has_skill_token("php", "laravel"):
        mismatch_hints.append("Stack PHP/Laravel trên CV khớp định hướng backend PHP.")

    # Stack hiện có lệch JD → kéo điểm match xuống
    stack_mismatch = bool(mismatch_hints) and any("lệch stack" in h.lower() for h in mismatch_hints)

    strengths = []
    if payload.skills:
        strengths.append(
            "CV đã liệt kê được nhóm kỹ năng rõ ràng: " + ", ".join(payload.skills[:6]) + "."
        )
    if re.search(r"project|dự án|portfolio|github", text, re.I):
        strengths.append("Có dấu hiệu dự án/portfolio — điểm cộng khi sàng CV fresher.")
    if not strengths:
        strengths.append("Hồ sơ đã có định hướng nghề, nhưng minh chứng còn mỏng.")

    weaknesses = list(mismatch_hints)
    if recommend.skill_gaps:
        weaknesses.append(
            "Thiếu hoặc chưa làm nổi các kỹ năng JD kỳ vọng: "
            + ", ".join(recommend.skill_gaps)
            + "."
        )
    if not re.search(r"\d+%|\d+\s*(users|người|api|tháng)", text, re.I):
        weaknesses.append("Mô tả chưa gắn số liệu/kết quả đo được — nhà tuyển dụng khó tin mức độ đóng góp.")

    improvements = [
        f"Viết lại phần tiêu đề/tóm tắt CV bám đúng vị trí {payload.target_job}.",
        "Mỗi dự án: bối cảnh → việc bạn làm → công nghệ → kết quả (số liệu).",
        "Đưa kỹ năng khớp JD lên đầu; kỹ năng lệch stack để nhóm phụ hoặc bỏ nếu không liên quan.",
        "Thêm link GitHub/demo nếu có sản phẩm thật.",
    ]

    salary_comment = None
    if payload.expected_salary and payload.expected_salary >= 20_000_000:
        salary_comment = (
            "Mức lương kỳ vọng khá cao so với hồ sơ fresher/sinh viên nếu chưa chứng minh được "
            "dự án mạnh và stack khớp JD — dễ bị loại ngay vòng CV."
        )

    score = recommend.match_score
    if stack_mismatch:
        score = min(score, 42)
        # Gợi ý học đúng stack JD thay vì giữ topic fullstack PHP
        recommend.skill_gaps = list(dict.fromkeys(
            ["Java", "Spring Boot", "MySQL", "REST API"] + list(recommend.skill_gaps)
        ))[:8]
        recommend.recommended_keyword_topics = ["Java", "Spring Boot", "Fullstack"]

    return EvaluateCVResponse(
        match_score=score,
        verdict="Cần chỉnh sửa CV trước khi nộp" if score < 70 else "Có thể đạt vòng CV nếu tinh chỉnh",
        overview=(
            f"Nhìn tổng thể, mức phù hợp với vị trí {payload.target_job} khoảng {score}/100. "
            + (mismatch_hints[0] + " " if mismatch_hints else "")
            + recommend.summary
        ),
        strengths=strengths[:5],
        weaknesses=weaknesses[:6],
        improvements=improvements,
        missing_items=recommend.skill_gaps[:5],
        skill_gaps=recommend.skill_gaps,
        interview_focus=[
            f"Vì sao ứng tuyển {payload.target_job} trong khi stack hiện tại trên CV là gì?",
            "Walkthrough 1 dự án sâu: kiến trúc, khó khăn, cách xử lý.",
        ],
        recommended_keyword_topics=recommend.recommended_keyword_topics,
        salary_comment=salary_comment,
    )


def _recommend_rules(payload: RecommendRequest) -> RecommendResponse:
    normalized_job = (payload.target_job or "").strip().lower()
    skill_lookup = {skill.lower(): skill for skill in payload.skills}

    market_skills = ["Git", "REST API", "Docker"]
    for keyword, skills in ROLE_SKILL_MAP.items():
        if keyword in normalized_job:
            market_skills = skills
            break

    current_skills_lower = set(skill_lookup.keys())
    gaps = [s for s in market_skills if s.lower() not in current_skills_lower]
    overlap = [s for s in market_skills if s.lower() in current_skills_lower]

    total = max(len(market_skills), 1)
    match_score = int(round((len(overlap) / total) * 100))
    match_score = max(35, min(match_score, 95))

    summary_parts: list[str] = []
    if overlap:
        summary_parts.append(
            f"Hồ sơ hiện tại đã có nền tảng phù hợp cho vị trí {payload.target_job}, "
            f"nổi bật ở các kỹ năng: {', '.join(overlap[:4])}."
        )
    else:
        summary_parts.append(
            f"Hồ sơ hiện tại mới ở mức khởi điểm so với vị trí {payload.target_job}, "
            "cần bổ sung thêm các kỹ năng lõi theo đúng vai trò mục tiêu."
        )

    if gaps:
        summary_parts.append(
            f"Các kỹ năng nên ưu tiên phát triển tiếp theo là: {', '.join(gaps)}."
        )

    if payload.cv_text:
        project_signals = re.search(
            r"project|dự án|portfolio", payload.cv_text, re.IGNORECASE
        )
        if not project_signals:
            summary_parts.append(
                "CV nên bổ sung thêm dự án thực tế hoặc portfolio "
                "để tăng sức thuyết phục khi ứng tuyển."
            )

    return RecommendResponse(
        match_score=match_score,
        skill_gaps=gaps,
        recommended_keyword_topics=gaps[:3] if gaps else overlap[:3],
        summary=" ".join(summary_parts).strip(),
    )


async def _recommend_with_ai(payload: RecommendRequest) -> RecommendResponse | None:
    skills_text = ", ".join(payload.skills) if payload.skills else "(chưa liệt kê)"
    cv_excerpt = (payload.cv_text or "")[:2500]

    messages = [
        {
            "role": "system",
            "content": (
                "Bạn là chuyên gia tư vấn nghề nghiệp cho sinh viên Việt Nam. "
                "Trả về JSON thuần (không markdown) với format:\n"
                "{\n"
                '  "match_score": 0-100,\n'
                '  "skill_gaps": ["kỹ năng thiếu 1", ...],\n'
                '  "recommended_keyword_topics": ["chủ đề học 1", ...],\n'
                '  "summary": "tóm tắt 2-4 câu tiếng Việt"\n'
                "}\n"
                "skill_gaps tối đa 6 mục, topics tối đa 5 mục."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Vị trí mục tiêu: {payload.target_job}\n"
                f"Kỹ năng hiện có: {skills_text}\n"
                f"Nội dung CV:\n{cv_excerpt or '(không có)'}\n\n"
                "Hãy đánh giá mức phù hợp và gợi ý kỹ năng cần bổ sung."
            ),
        },
    ]

    raw = await complete_json(
        provider=payload.provider or "claude",
        api_key=payload.api_key or "",
        model=payload.model,
        messages=messages,
        short_messages=[
            {
                "role": "system",
                "content": (
                    'Chỉ trả JSON: {"match_score":0,"skill_gaps":[],'
                    '"recommended_keyword_topics":[],"summary":"..."}. Không markdown.'
                ),
            },
            {
                "role": "user",
                "content": (
                    f"Vị trí: {payload.target_job}. Kỹ năng: {skills_text}. "
                    f"CV rút: {(payload.cv_text or '')[:800] or '(trống)'}. JSON ngắn."
                ),
            },
        ],
        temperature=0.4,
        is_valid=lambda d: bool(str(d.get("summary") or "").strip()),
    )

    try:
        score = int(raw.get("match_score") or 0)
        score = max(0, min(score, 100))
        gaps = [str(x) for x in (raw.get("skill_gaps") or []) if x][:6]
        topics = [str(x) for x in (raw.get("recommended_keyword_topics") or []) if x][:5]
        summary = str(raw.get("summary") or "").strip()
        return RecommendResponse(
            match_score=score,
            skill_gaps=gaps,
            recommended_keyword_topics=topics or gaps[:3],
            summary=summary,
        )
    except (TypeError, ValueError) as exc:
        logger.warning("Failed to map career AI JSON: %s", exc)
        return None
